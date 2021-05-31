import csv

try:
    from nltk import wordpunct_tokenize
    from nltk.corpus import stopwords
except ImportError:
    print (' You need to install nltk (http://nltk.org/index.html)')


def calculate_languages_ratios(text):
    languages_ratios = {}
    tokens = wordpunct_tokenize(text)
    words = [word.lower() for word in tokens]

    # Compute per language included in nltk number of unique stopwords appearing in analyzed text
    for language in stopwords.fileids():
        stopwords_set = set(stopwords.words(language))
        words_set = set(words)
        common_elements = words_set.intersection(stopwords_set)

        languages_ratios[language] = len(common_elements) # language "score"

    return languages_ratios

def detect_language(text):
    ratios = calculate_languages_ratios(text)
    most_rated_language = max(ratios, key=ratios.get)
    return most_rated_language

def load_bad_words(language):
    data = []
    if language.upper() in ['ENGLISH','FRENCH','SPANISH','GERMAN']:
        f= open("../public/datasets/"+language.lower()+'.csv')
        reader = csv.reader(f)
        for row in reader:
            data += [row[0].lower()]

        return  data

def load_file(filename):
	file=open(filename,'rb')
	return file

#pdf 
def getTextPdf(path):
    from io import StringIO
    from pdfminer.converter import TextConverter
    from pdfminer.layout import LAParams
    from pdfminer.pdfdocument import PDFDocument
    from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
    from pdfminer.pdfpage import PDFPage
    from pdfminer.pdfparser import PDFParser

    output_string = StringIO()
    with open(path, 'rb') as in_file:
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)
    return  output_string.getvalue()

#Docs
def getTextDOCX(filename):

    from os import path
    import docx

    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

